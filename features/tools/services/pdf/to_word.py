"""PDF to Word conversion."""
import fitz
from docx import Document

from features.tools.models import ToolJob
from features.tools.services import file_handler


def _max_pages(job: ToolJob) -> int:
    if job.user and job.user.is_authenticated:
        try:
            profile = job.user.profile
            if profile.is_on_trial or (profile.plan and profile.plan.slug in ("pro", "business")):
                return 9999
        except Exception:
            pass
    return 1


def run(job: ToolJob) -> dict:
    src_rel = job.input_files[0]
    src = file_handler.to_absolute(src_rel)
    pdf = fitz.open(src)
    page_count = pdf.page_count
    max_pages = _max_pages(job)

    document = Document()
    for i in range(min(max_pages, page_count)):
        text = pdf[i].get_text()
        if text.strip():
            document.add_paragraph(text)
        else:
            document.add_paragraph(f"[Page {i + 1} — no extractable text]")

    if page_count > max_pages:
        document.add_paragraph(
            "\n[Preview: first page only. Upgrade to Pro for full document conversion.]"
        )
    pdf.close()

    out = file_handler.output_path(job, "converted.docx")
    document.save(str(out))
    out_rel = file_handler.to_relative(out)
    return {
        "output_file": out_rel,
        "compression_stat": {"pages": min(max_pages, page_count)},
        "preview_file": "",
    }
